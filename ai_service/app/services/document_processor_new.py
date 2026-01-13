"""
Comprehensive Document Processing Service
Handles text extraction, OCR, embeddings, and ChromaDB integration
"""

import io
import os
from typing import Tuple, List, Dict, Optional
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from pdf2image import convert_from_bytes
import pytesseract
from app.core.chroma_client import chroma_manager
from app.utils.logger import logger


class ComprehensiveDocumentProcessor:
    """Enhanced document processor with OCR and embedding capabilities"""
    
    def __init__(self):
        """Initialize the document processor"""
        self.embedding_model = None
        self.chroma_collection = None
        
    def _ensure_embedding_model(self):
        """Lazy load embedding model"""
        if self.embedding_model is None:
            try:
                from sentence_transformers import SentenceTransformer
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                logger.info("Embedding model loaded: all-MiniLM-L6-v2")
            except Exception as e:
                logger.error(f"Failed to load embedding model: {e}")
                raise
    
    def _ensure_chroma_collection(self):
        """Lazy load ChromaDB collection"""
        if self.chroma_collection is None:
            try:
                self.chroma_collection = chroma_manager.get_or_create_collection(
                    name="user_documents",
                    metadata={"description": "User uploaded documents with embeddings"}
                )
                logger.info("ChromaDB collection initialized: user_documents")
            except Exception as e:
                logger.error(f"Failed to initialize ChromaDB collection: {e}")
                raise
    
    async def extract_text_from_pdf(self, content: bytes) -> Tuple[str, bool]:
        """
        Extract text from PDF bytes
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Tuple of (extracted_text, needs_ocr)
        """
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            text = ""
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            
            # Check if text extraction was successful
            word_count = len(text.split())
            needs_ocr = word_count < 50  # If less than 50 words, likely scanned
            
            logger.info(f"Extracted {len(text)} characters from PDF (words: {word_count}, needs_ocr: {needs_ocr})")
            return text.strip(), needs_ocr
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    async def extract_text_from_word(self, content: bytes) -> str:
        """
        Extract text from Word document bytes
        
        Args:
            content: Word file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from Word document")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from Word: {e}")
            raise
    
    async def perform_ocr_on_pdf(self, content: bytes, language: str = 'eng') -> str:
        """
        Perform OCR on scanned PDF
        
        Args:
            content: PDF file content as bytes
            language: OCR language (default: eng)
            
        Returns:
            Extracted text from OCR
        """
        try:
            # Convert PDF pages to images
            images = convert_from_bytes(content, dpi=300)
            text = ""
            
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image, lang=language)
                text += f"\n--- Page {i+1} ---\n{page_text}"
            
            logger.info(f"OCR extracted {len(text)} characters from {len(images)} pages")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error performing OCR on PDF: {e}")
            raise
    
    async def perform_ocr_on_image(self, content: bytes, language: str = 'eng') -> str:
        """
        Perform OCR on image
        
        Args:
            content: Image file content as bytes
            language: OCR language (default: eng)
            
        Returns:
            Extracted text from OCR
        """
        try:
            from PIL import Image
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image, lang=language)
            
            logger.info(f"OCR extracted {len(text)} characters from image")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error performing OCR on image: {e}")
            raise
    
    def chunk_text(self, text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to chunk
            chunk_size: Number of words per chunk
            chunk_overlap: Number of words to overlap
            
        Returns:
            List of text chunks
        """
        words = text.split()
        chunks = []
        
        i = 0
        while i < len(words):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            chunks.append(chunk_text)
            i += chunk_size - chunk_overlap
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
    
    async def generate_embeddings(
        self,
        file_id: str,
        user_id: str,
        text: str,
        metadata: Dict,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> Dict:
        """
        Generate embeddings and store in ChromaDB
        
        Args:
            file_id: Unique file identifier
            user_id: User identifier
            text: Text to embed
            metadata: Additional metadata
            chunk_size: Words per chunk
            chunk_overlap: Overlap between chunks
            
        Returns:
            Dict with embedding IDs and chunk count
        """
        try:
            self._ensure_embedding_model()
            self._ensure_chroma_collection()
            
            # Chunk the text
            chunks = self.chunk_text(text, chunk_size, chunk_overlap)
            
            # Generate embeddings
            embeddings = self.embedding_model.encode(chunks).tolist()
            
            # Prepare IDs and metadata
            chunk_ids = [f"{file_id}_chunk_{i}" for i in range(len(chunks))]
            chunk_metadata = []
            for i, chunk in enumerate(chunks):
                chunk_meta = {
                    "file_id": file_id,
                    "user_id": user_id,
                    "chunk_index": i,
                    **metadata
                }
                chunk_metadata.append(chunk_meta)
            
            # Store in ChromaDB
            self.chroma_collection.add(
                ids=chunk_ids,
                embeddings=embeddings,
                documents=chunks,
                metadatas=chunk_metadata
            )
            
            logger.info(f"Generated {len(chunks)} embeddings for file {file_id}")
            
            return {
                "success": True,
                "chunk_count": len(chunks),
                "embedding_ids": chunk_ids,
                "collection": "user_documents"
            }
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise
    
    async def search_embeddings(
        self,
        user_id: str,
        query: str,
        limit: int = 10,
        filters: Optional[Dict] = None
    ) -> List[Dict]:
        """
        Search documents using semantic similarity
        
        Args:
            user_id: User identifier
            query: Search query
            limit: Max results
            filters: Additional filters
            
        Returns:
            List of matching documents with scores
        """
        try:
            self._ensure_embedding_model()
            self._ensure_chroma_collection()
            
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0].tolist()
            
            # Prepare where clause
            where_clause = {"user_id": user_id}
            if filters:
                where_clause.update(filters)
            
            # Search ChromaDB
            results = self.chroma_collection.query(
                query_embeddings=[query_embedding],
                n_results=limit,
                where=where_clause
            )
            
            # Format results
            formatted_results = []
            if results and results.get('ids') and len(results['ids']) > 0:
                for i in range(len(results['ids'][0])):
                    formatted_results.append({
                        "id": results['ids'][0][i],
                        "document": results['documents'][0][i] if results.get('documents') else None,
                        "metadata": results['metadatas'][0][i] if results.get('metadatas') else {},
                        "distance": results['distances'][0][i] if results.get('distances') else None,
                        "score": 1 - results['distances'][0][i] if results.get('distances') else None,
                    })
            
            logger.info(f"Found {len(formatted_results)} results for query: {query}")
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error searching embeddings: {e}")
            raise
    
    async def delete_embeddings(self, ids: List[str]) -> bool:
        """
        Delete embeddings from ChromaDB
        
        Args:
            ids: List of embedding IDs to delete
            
        Returns:
            True if successful
        """
        try:
            self._ensure_chroma_collection()
            
            self.chroma_collection.delete(ids=ids)
            
            logger.info(f"Deleted {len(ids)} embeddings")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting embeddings: {e}")
            return False


# Create singleton instance with new name to avoid conflicts
comprehensive_document_processor = ComprehensiveDocumentProcessor()
