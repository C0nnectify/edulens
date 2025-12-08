"""
SOP Export Service for document generation.

This service handles exporting SOP drafts to various formats including
DOCX and PDF with proper formatting and styling.
"""

import io
import logging
from typing import BinaryIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from ..models.sop_generator import SOPDraft

logger = logging.getLogger(__name__)


class SOPExportService:
    """
    Service for exporting SOP drafts to various formats.

    Supports DOCX and PDF export with professional formatting.
    """

    @staticmethod
    def export_to_txt(draft: SOPDraft) -> bytes:
        """
        Export SOP draft to plain text.

        Args:
            draft: SOP draft to export

        Returns:
            Plain text content as bytes
        """
        content = f"""STATEMENT OF PURPOSE

{draft.content}

---
Generated: {draft.generated_at.strftime('%Y-%m-%d %H:%M')}
Tone: {draft.tone.value}
Word Count: {draft.word_count}
"""
        return content.encode('utf-8')

    @staticmethod
    def export_to_docx(
        draft: SOPDraft,
        applicant_name: str = None,
        program_name: str = None,
        university_name: str = None,
    ) -> bytes:
        """
        Export SOP draft to DOCX format with professional formatting.

        Args:
            draft: SOP draft to export
            applicant_name: Name of the applicant (optional)
            program_name: Name of the program (optional)
            university_name: Name of the university (optional)

        Returns:
            DOCX file content as bytes
        """
        # Create document
        doc = Document()

        # Set up styles
        styles = doc.styles

        # Configure normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)

        # Create title style if it doesn't exist
        try:
            title_style = styles['Title']
        except KeyError:
            title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)

        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(16)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add header information if provided
        if applicant_name or program_name or university_name:
            header_paragraph = doc.add_paragraph()
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if applicant_name:
                run = header_paragraph.add_run(applicant_name)
                run.bold = True
                run.font.size = Pt(14)
                header_paragraph.add_run('\n')

            if program_name and university_name:
                header_paragraph.add_run(f"{program_name}\n{university_name}")
            elif program_name:
                header_paragraph.add_run(program_name)
            elif university_name:
                header_paragraph.add_run(university_name)

            doc.add_paragraph()  # Empty line

        # Add title
        title = doc.add_heading('Statement of Purpose', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add empty line
        doc.add_paragraph()

        # Add content paragraphs
        paragraphs = draft.content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                paragraph = doc.add_paragraph(para_text.strip())
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.5
                paragraph_format.space_after = Pt(12)
                paragraph_format.first_line_indent = Inches(0)
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add footer with generation info (optional, commented out by default)
        # doc.add_page_break()
        # footer = doc.add_paragraph()
        # footer.add_run(f"Generated: {draft.generated_at.strftime('%B %d, %Y')}\n")
        # footer.add_run(f"Word Count: {draft.word_count}")
        # footer_format = footer.paragraph_format
        # footer_format.space_before = Pt(12)
        # footer.style = 'Normal'
        # for run in footer.runs:
        #     run.font.size = Pt(9)
        #     run.font.color.rgb = RGBColor(128, 128, 128)

        # Save to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)

        return docx_file.getvalue()

    @staticmethod
    def export_to_pdf(
        draft: SOPDraft,
        applicant_name: str = None,
        program_name: str = None,
        university_name: str = None,
    ) -> bytes:
        """
        Export SOP draft to PDF format with professional formatting.

        Args:
            draft: SOP draft to export
            applicant_name: Name of the applicant (optional)
            program_name: Name of the program (optional)
            university_name: Name of the university (optional)

        Returns:
            PDF file content as bytes
        """
        # Create PDF buffer
        pdf_buffer = io.BytesIO()

        # Create document
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            topMargin=1*inch,
            bottomMargin=1*inch,
            leftMargin=1*inch,
            rightMargin=1*inch,
        )

        # Container for flowables
        story = []

        # Get styles
        styles = getSampleStyleSheet()

        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='black',
            spaceAfter=12,
            alignment=1,  # Center
            fontName='Times-Bold',
        )

        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Normal'],
            fontSize=12,
            textColor='black',
            spaceAfter=6,
            alignment=1,  # Center
            fontName='Times-Roman',
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=12,
            textColor='black',
            spaceAfter=12,
            alignment=4,  # Justify
            fontName='Times-Roman',
            leading=18,  # Line spacing (1.5 * fontSize)
            firstLineIndent=0,
        )

        footer_style = ParagraphStyle(
            'CustomFooter',
            parent=styles['Normal'],
            fontSize=9,
            textColor='gray',
            spaceAfter=6,
            fontName='Times-Roman',
        )

        # Add header information if provided
        if applicant_name:
            story.append(Paragraph(f"<b>{applicant_name}</b>", header_style))

        if program_name and university_name:
            story.append(Paragraph(f"{program_name}", header_style))
            story.append(Paragraph(f"{university_name}", header_style))
        elif program_name:
            story.append(Paragraph(program_name, header_style))
        elif university_name:
            story.append(Paragraph(university_name, header_style))

        if applicant_name or program_name or university_name:
            story.append(Spacer(1, 0.2*inch))

        # Add title
        story.append(Paragraph("<b>Statement of Purpose</b>", title_style))
        story.append(Spacer(1, 0.3*inch))

        # Add content paragraphs
        paragraphs = draft.content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Clean and format text for PDF
                cleaned_text = para_text.strip().replace('\n', ' ')
                story.append(Paragraph(cleaned_text, body_style))

        # Build PDF
        doc.build(story)

        # Get PDF bytes
        pdf_buffer.seek(0)
        return pdf_buffer.getvalue()

    @staticmethod
    def export_draft(
        draft: SOPDraft,
        format: str,
        applicant_name: str = None,
        program_name: str = None,
        university_name: str = None,
    ) -> tuple[bytes, str, str]:
        """
        Export draft in the specified format.

        Args:
            draft: SOP draft to export
            format: Export format (txt, docx, pdf)
            applicant_name: Name of the applicant (optional)
            program_name: Name of the program (optional)
            university_name: Name of the university (optional)

        Returns:
            Tuple of (file_content, media_type, filename)

        Raises:
            ValueError: If format is unsupported
        """
        format = format.lower()

        if format == "txt":
            content = SOPExportService.export_to_txt(draft)
            media_type = "text/plain"
            filename = f"statement_of_purpose_{draft.draft_id[:8]}.txt"

        elif format == "docx":
            content = SOPExportService.export_to_docx(
                draft,
                applicant_name,
                program_name,
                university_name
            )
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"statement_of_purpose_{draft.draft_id[:8]}.docx"

        elif format == "pdf":
            content = SOPExportService.export_to_pdf(
                draft,
                applicant_name,
                program_name,
                university_name
            )
            media_type = "application/pdf"
            filename = f"statement_of_purpose_{draft.draft_id[:8]}.pdf"

        else:
            raise ValueError(f"Unsupported format: {format}. Use txt, docx, or pdf.")

        logger.info(f"Exported draft {draft.draft_id} to {format}")

        return content, media_type, filename
