"""
Document processing service for extracting text from various file formats
"""

import io
from typing import Tuple
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from app.utils.logger import logger
from app.utils.file_utils import get_file_extension


class DocumentProcessor:
    """Process documents and extract text content"""

    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            reader = PdfReader(file_path)
            text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

            logger.info(f"Extracted {len(text)} characters from PDF: {file_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

    @staticmethod
    async def extract_text_from_txt(file_path: str) -> str:
        """
        Extract text from TXT file

        Args:
            file_path: Path to TXT file

        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

            logger.info(f"Extracted {len(text)} characters from TXT: {file_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise

    async def process_document(self, file_path: str) -> Tuple[str, str]:
        """
        Process document and extract text based on file type

        Args:
            file_path: Path to document file

        Returns:
            Tuple of (extracted_text, file_type)
        """
        file_path_obj = Path(file_path)
        extension = get_file_extension(file_path_obj.name)

        if extension == 'pdf':
            text = await self.extract_text_from_pdf(file_path)
            file_type = 'pdf'
        elif extension in ['doc', 'docx']:
            text = await self.extract_text_from_docx(file_path)
            file_type = 'docx'
        elif extension == 'txt':
            text = await self.extract_text_from_txt(file_path)
            file_type = 'txt'
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        return text, file_type

    @staticmethod
    def validate_extracted_text(text: str, min_length: int = 10) -> bool:
        """
        Validate that extracted text meets minimum requirements

        Args:
            text: Extracted text
            min_length: Minimum text length

        Returns:
            True if text is valid
        """
        if not text or len(text.strip()) < min_length:
            logger.warning(f"Extracted text too short: {len(text)} characters")
            return False

        return True


# Create singleton instance
document_processor = DocumentProcessor()
